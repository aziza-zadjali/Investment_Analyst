"""
File processing utilities for PDF, DOCX, and XLSX files
"""

import io
import fitz  # PyMuPDF
import docx
import openpyxl
import pandas as pd
from typing import Dict, List, Any, Optional
import streamlit as st

class FileProcessor:
    """Handle document processing for various file types"""
    
    @staticmethod
    def process_pdf(file) -> Dict[str, Any]:
        """Extract text and metadata from PDF files"""
        try:
            # Read PDF
            pdf_bytes = file.read()
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            
            # Extract text from all pages
            text_content = []
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                text_content.append(page.get_text())
            
            # Extract metadata
            metadata = pdf_document.metadata
            
            result = {
                'text': '\n'.join(text_content),
                'pages': pdf_document.page_count,
                'metadata': metadata,
                'type': 'pdf'
            }
            
            pdf_document.close()
            return result
            
        except Exception as e:
            st.error(f"Error processing PDF: {str(e)}")
            return None
    
    @staticmethod
    def process_docx(file) -> Dict[str, Any]:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file)
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            result = {
                'text': '\n'.join(paragraphs),
                'paragraphs': paragraphs,
                'tables': tables,
                'num_tables': len(tables),
                'type': 'docx'
            }
            
            return result
            
        except Exception as e:
            st.error(f"Error processing DOCX: {str(e)}")
            return None
    
    @staticmethod
    def process_xlsx(file) -> Dict[str, Any]:
        """Extract data from Excel files"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file)
            sheets_data = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file, sheet_name=sheet_name)
                sheets_data[sheet_name] = df
            
            # Generate summary
            summary = {
                'num_sheets': len(excel_file.sheet_names),
                'sheet_names': excel_file.sheet_names,
                'sheets_data': sheets_data
            }
            
            # Convert to text representation
            text_parts = []
            for sheet_name, df in sheets_data.items():
                text_parts.append(f"\n=== Sheet: {sheet_name} ===\n")
                text_parts.append(df.to_string())
            
            result = {
                'text': '\n'.join(text_parts),
                'summary': summary,
                'dataframes': sheets_data,
                'type': 'xlsx'
            }
            
            return result
            
        except Exception as e:
            st.error(f"Error processing XLSX: {str(e)}")
            return None
    
    @staticmethod
    def process_txt(file) -> Dict[str, Any]:
        """Extract text from TXT files"""
        try:
            text = file.read().decode('utf-8')
            
            result = {
                'text': text,
                'lines': text.split('\n'),
                'type': 'txt'
            }
            
            return result
            
        except Exception as e:
            st.error(f"Error processing TXT: {str(e)}")
            return None
    
    @classmethod
    def process_file(cls, file) -> Optional[Dict[str, Any]]:
        """Process file based on extension"""
        file_extension = file.name.split('.')[-1].lower()
        
        processors = {
            'pdf': cls.process_pdf,
            'docx': cls.process_docx,
            'xlsx': cls.process_xlsx,
            'txt': cls.process_txt
        }
        
        processor = processors.get(file_extension)
        if processor:
            return processor(file)
        else:
            st.error(f"Unsupported file type: {file_extension}")
            return None
    
    @staticmethod
    def extract_financial_data(text: str) -> Dict[str, Any]:
        """Extract financial metrics from text using pattern matching"""
        import re
        
        metrics = {}
        
        # Revenue patterns
        revenue_patterns = [
            r'revenue[:\s]+\$?([\d,]+\.?\d*)\s*(million|M|billion|B)?',
            r'sales[:\s]+\$?([\d,]+\.?\d*)\s*(million|M|billion|B)?',
        ]
        
        # EBITDA patterns
        ebitda_patterns = [
            r'EBITDA[:\s]+\$?([\d,]+\.?\d*)\s*(million|M|billion|B)?',
            r'ebitda[:\s]+\$?([\d,]+\.?\d*)\s*(million|M|billion|B)?',
        ]
        
        # Extract metrics
        for pattern in revenue_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                metrics['revenue'] = match.group(1)
                metrics['revenue_unit'] = match.group(2) if match.group(2) else 'unknown'
                break
        
        for pattern in ebitda_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                metrics['ebitda'] = match.group(1)
                metrics['ebitda_unit'] = match.group(2) if match.group(2) else 'unknown'
                break
        
        return metrics
